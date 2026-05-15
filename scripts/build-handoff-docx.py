#!/usr/bin/env python3
"""
Convert HANDOFF.md into a polished Maryland_Budget_Intel_Tool_Handoff.docx
that can be emailed to a client. Re-run any time HANDOFF.md changes.

Usage:
    pip install python-docx
    python scripts/build-handoff-docx.py
"""

from __future__ import annotations
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "HANDOFF.md"
DST = ROOT / "docs" / "Maryland_Budget_Intel_Tool_Handoff.docx"

PURPLE = RGBColor(0x63, 0x21, 0xA5)
DEEP   = RGBColor(0x21, 0x10, 0x30)
INK    = RGBColor(0x12, 0x12, 0x12)
GREY   = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0x80, 0x2C, 0xD7)


def set_font(run, *, size=11, bold=False, italic=False, color=INK, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    sizes = {1: 22, 2: 16, 3: 13}
    colors = {1: DEEP, 2: PURPLE, 3: ACCENT}
    run = p.add_run(text)
    set_font(run, size=sizes.get(level, 11), bold=True, color=colors.get(level, INK))
    return p


def add_paragraph_runs(doc, text):
    """Render a markdown paragraph with **bold**, _italic_, `code`, and [link](url) handled inline."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    _render_inline(p, text)
    return p


_INLINE_RX = re.compile(
    r"(\*\*[^*]+\*\*"                 # bold
    r"|`[^`]+`"                       # code
    r"|\[[^\]]+\]\([^)]+\)"           # links
    r"|_[^_]+_)"                      # italic
)


def _render_inline(paragraph, text):
    pos = 0
    for m in _INLINE_RX.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.font.bold = True
        elif token.startswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
            r.font.color.rgb = ACCENT
            r.font.size = Pt(10)
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            r = paragraph.add_run(label)
            r.font.color.rgb = PURPLE
            r.font.underline = True
        elif token.startswith("_"):
            r = paragraph.add_run(token[1:-1])
            r.font.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_bullet(doc, text, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    if indent:
        p.paragraph_format.left_indent = Inches(0.25 * indent)
    p.paragraph_format.space_after = Pt(3)
    _render_inline(p, text)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    r.font.color.rgb = DEEP


def add_table_from_md(doc, header_cells, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
    table.style = "Light Grid Accent 4"
    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(header_cells):
        cell = hdr[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=10, bold=True, color=DEEP)
    # Body
    for ri, row in enumerate(rows, start=1):
        for ci, cell_text in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            _render_inline(p, cell_text)
            for run in p.runs:
                run.font.size = Pt(10)
                if run.font.color.rgb is None:
                    run.font.color.rgb = INK
    doc.add_paragraph()  # space after


def add_screenshot_placeholder(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"[ SCREENSHOT PLACEHOLDER ]\n{caption}")
    set_font(r, size=10, italic=True, color=GREY)
    # border-style hint via a horizontal rule paragraph
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sep.add_run("─" * 60)
    set_font(sr, size=9, color=GREY)


def render_cover(doc):
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("Maryland Budget Intel Tool")
    set_font(r, size=32, bold=True, color=DEEP)
    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Technical Handoff Documentation")
    set_font(r, size=18, color=PURPLE)
    # Affiliation
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("University of Maryland iSchool Capstone Project")
    set_font(r, size=12, italic=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("In partnership with NXT and the Senate Budget & Taxation Committee Staff")
    set_font(r, size=11, italic=True, color=GREY)

    # Spacer
    for _ in range(2):
        doc.add_paragraph()

    # Meta box
    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.style = "Light Shading Accent 4"
    rows = [
        ("Live Site", "https://mdbudgetintel.netlify.app"),
        ("Home Page Repository", "https://github.com/asingh3506/md-budget-site"),
        ("Deep-Dive Repository", "https://github.com/priyanshu124/IT_state_budget_agent"),
        ("Document Version", "1.0 — May 2026"),
    ]
    for i, (k, v) in enumerate(rows):
        c1, c2 = meta.rows[i].cells
        c1.text = ""
        c2.text = ""
        r1 = c1.paragraphs[0].add_run(k)
        set_font(r1, size=11, bold=True, color=DEEP)
        r2 = c2.paragraphs[0].add_run(v)
        set_font(r2, size=11, color=INK)

    # Contributors
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("Aarushi Singh  ·  Nadvi Haque  ·  Priyanshu Gupta  ·  James Van Doorn")
    set_font(r, size=11, color=INK)

    doc.add_page_break()


def parse_md(text):
    """Parse the simplified subset of markdown we use. Yields (kind, payload)."""
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.rstrip()
        if not s.strip():
            i += 1
            continue
        # Headings
        m = re.match(r"^(#{1,3})\s+(.*)$", s)
        if m:
            level = len(m.group(1))
            yield ("heading", level, m.group(2))
            i += 1
            continue
        # Horizontal rule
        if re.match(r"^---+\s*$", s):
            yield ("hr",)
            i += 1
            continue
        # Code block
        if s.startswith("```"):
            buf = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            yield ("code", "\n".join(buf))
            continue
        # Table
        if "|" in s and i + 1 < len(lines) and re.match(r"^\s*\|?[-:\s|]+\|?\s*$", lines[i+1]):
            header = [c.strip() for c in s.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            yield ("table", header, rows)
            continue
        # Blockquote
        if s.startswith("> "):
            yield ("quote", s[2:].strip())
            i += 1
            continue
        # Screenshot placeholder pattern: `[SCREENSHOT: ...]`
        m = re.match(r"^`\[SCREENSHOT:\s*(.+?)\]`\s*$", s.strip())
        if m:
            yield ("screenshot", m.group(1))
            i += 1
            continue
        # Bullet
        m = re.match(r"^(\s*)[-*]\s+(.*)$", s)
        if m:
            indent = len(m.group(1)) // 2
            yield ("bullet", indent, m.group(2))
            i += 1
            continue
        # Italics-only line (notes etc.)
        if s.startswith("*") and s.endswith("*") and not s.startswith("**"):
            yield ("italic", s.strip("*").strip())
            i += 1
            continue
        # Paragraph (collect continuation lines until blank or block element)
        para = [s]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if not nxt.strip(): break
            if re.match(r"^#{1,3}\s|^---+$|^```|^>\s|^\s*[-*]\s|^\|", nxt): break
            para.append(nxt)
            i += 1
        yield ("paragraph", " ".join(para))


def main():
    md = SRC.read_text()
    # Strip the markdown YAML/title block we'll render ourselves on the cover page
    md = re.sub(r"^# Maryland Budget Intel Tool\s*\n.*?(?=\n## Table of Contents)", "",
                md, count=1, flags=re.DOTALL)

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    render_cover(doc)

    for node in parse_md(md):
        kind = node[0]
        if kind == "heading":
            add_heading(doc, node[2], node[1])
        elif kind == "paragraph":
            add_paragraph_runs(doc, node[1])
        elif kind == "bullet":
            add_bullet(doc, node[2], indent=node[1])
        elif kind == "code":
            add_code_block(doc, node[1])
        elif kind == "table":
            add_table_from_md(doc, node[1], node[2])
        elif kind == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(node[1])
            set_font(r, size=11, italic=True, color=GREY)
        elif kind == "screenshot":
            add_screenshot_placeholder(doc, node[1])
        elif kind == "italic":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(node[1])
            set_font(r, size=10, italic=True, color=GREY)
        elif kind == "hr":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("─" * 50)
            set_font(r, size=9, color=GREY)

    doc.save(DST)
    print(f"✓ Wrote {DST.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
